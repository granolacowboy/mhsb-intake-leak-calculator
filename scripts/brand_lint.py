#!/usr/bin/env python3
"""brand_lint.py - config-driven brand compliance lint gate.

Scans text-like files for banned phrases, banned/legacy hex colors,
retired fonts, and punctuation violations, as defined by a JSON brand
config. Designed to run as a ship gate: exit 1 on any error-level
finding.

Usage:
    python brand_lint.py <brand-config.json> <path> [<path> ...]
        [--ext .md,.html,...] [--json] [--warnings-as-errors] [--docx]

Exit codes:
    0  clean (or warnings only, unless --warnings-as-errors)
    1  one or more error-level findings
    2  usage or configuration problem

Config schema (all keys optional unless noted):
{
  "brand": "Brand Name",              # informational
  "version": "4.3",                   # informational, printed in header
  "banned_phrases": ["..."],          # error; case-insensitive substring
  "banned_regexes": ["\\b..."],       # error; case-insensitive regex
  "warn_phrases": ["..."],            # warning; case-insensitive substring
  "warn_regexes": ["\\b..."],         # warning; case-insensitive regex
  "banned_hexes": ["#1F4788"],        # error; '#' optional, case-insensitive
  "retired_fonts": ["Fraunces"],      # error; case-insensitive substring
  "punctuation": {
      "forbid_em_dash": true,         # error on U+2014
      "forbid_en_dash": true          # error on U+2013
  }
}
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

DEFAULT_EXTENSIONS: frozenset[str] = frozenset(
    {
        ".md", ".markdown", ".txt", ".html", ".htm", ".css", ".scss",
        ".js", ".jsx", ".ts", ".tsx", ".py", ".json", ".svg", ".xml",
        ".yml", ".yaml", ".csv", ".toml", ".mdx",
    }
)
EXCLUDED_DIRS: frozenset[str] = frozenset(
    {".git", "node_modules", "__pycache__", "dist", "build", ".venv", "venv"}
)
MAX_FILE_BYTES: int = 5 * 1024 * 1024

# Bare 6-digit hex only in obvious code/color contexts to limit noise:
# quoted ("384152"), or preceded by '#', or a key like "primary": 384152.
HEX_WITH_HASH_RE = re.compile(r"#([0-9a-fA-F]{6}|[0-9a-fA-F]{3})\b")
HEX_QUOTED_RE = re.compile(r"['\"]([0-9a-fA-F]{6})['\"]")

EM_DASH = "—"
EN_DASH = "–"


@dataclass
class Finding:
    severity: str  # "error" | "warning"
    path: str
    line: int
    rule: str
    detail: str

    def to_dict(self) -> dict[str, object]:
        return {
            "severity": self.severity,
            "path": self.path,
            "line": self.line,
            "rule": self.rule,
            "detail": self.detail,
        }


@dataclass
class BrandConfig:
    brand: str = "unnamed brand"
    version: str = "unversioned"
    banned_phrases: list[str] = field(default_factory=list)
    banned_regexes: list[re.Pattern[str]] = field(default_factory=list)
    warn_phrases: list[str] = field(default_factory=list)
    warn_regexes: list[re.Pattern[str]] = field(default_factory=list)
    banned_hexes: set[str] = field(default_factory=set)  # normalized, no '#', lower
    retired_fonts: list[str] = field(default_factory=list)
    forbid_em_dash: bool = False
    forbid_en_dash: bool = False

    @classmethod
    def load(cls, path: Path) -> "BrandConfig":
        try:
            raw = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError) as exc:
            raise SystemExit(f"config error: cannot read {path}: {exc}") from exc
        if not isinstance(raw, dict):
            raise SystemExit(f"config error: {path} must contain a JSON object")

        def _compile(patterns: list[str], key: str) -> list[re.Pattern[str]]:
            compiled: list[re.Pattern[str]] = []
            for pat in patterns:
                try:
                    compiled.append(re.compile(pat, re.IGNORECASE))
                except re.error as exc:
                    raise SystemExit(
                        f"config error: bad regex in {key!r}: {pat!r} ({exc})"
                    ) from exc
            return compiled

        punct = raw.get("punctuation", {}) or {}
        return cls(
            brand=str(raw.get("brand", "unnamed brand")),
            version=str(raw.get("version", "unversioned")),
            banned_phrases=[str(p) for p in raw.get("banned_phrases", [])],
            banned_regexes=_compile(
                [str(p) for p in raw.get("banned_regexes", [])], "banned_regexes"
            ),
            warn_phrases=[str(p) for p in raw.get("warn_phrases", [])],
            warn_regexes=_compile(
                [str(p) for p in raw.get("warn_regexes", [])], "warn_regexes"
            ),
            banned_hexes={
                str(h).lstrip("#").lower() for h in raw.get("banned_hexes", [])
            },
            retired_fonts=[str(f) for f in raw.get("retired_fonts", [])],
            forbid_em_dash=bool(punct.get("forbid_em_dash", False)),
            forbid_en_dash=bool(punct.get("forbid_en_dash", False)),
        )


def iter_target_files(paths: list[Path], extensions: frozenset[str]) -> Iterator[Path]:
    seen: set[Path] = set()
    for base in paths:
        if base.is_file():
            candidates: Iterator[Path] = iter([base])
        elif base.is_dir():
            candidates = (
                p
                for p in sorted(base.rglob("*"))
                if p.is_file() and not (EXCLUDED_DIRS & set(p.parts))
            )
        else:
            print(f"warning: path not found, skipping: {base}", file=sys.stderr)
            continue
        for path in candidates:
            resolved = path.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)
            if path.suffix.lower() in extensions:
                yield path


def _expand_hexes(hexes: set[str]) -> set[str]:
    """Expand 3-digit shorthand to 6-digit for comparison."""
    expanded = set(hexes)
    for h in hexes:
        if len(h) == 3:
            expanded.add("".join(c * 2 for c in h))
    return expanded


def lint_text(path: Path, text: str, cfg: BrandConfig) -> list[Finding]:
    findings: list[Finding] = []
    rel = str(path)
    banned_hexes = _expand_hexes(cfg.banned_hexes)

    for lineno, line in enumerate(text.splitlines(), start=1):
        lowered = line.lower()

        for phrase in cfg.banned_phrases:
            if phrase.lower() in lowered:
                findings.append(
                    Finding("error", rel, lineno, "banned-phrase", phrase)
                )
        for phrase in cfg.warn_phrases:
            if phrase.lower() in lowered:
                findings.append(
                    Finding("warning", rel, lineno, "avoid-phrase", phrase)
                )
        for pattern in cfg.banned_regexes:
            if pattern.search(line):
                findings.append(
                    Finding("error", rel, lineno, "banned-pattern", pattern.pattern)
                )
        for pattern in cfg.warn_regexes:
            if pattern.search(line):
                findings.append(
                    Finding("warning", rel, lineno, "avoid-pattern", pattern.pattern)
                )
        for font in cfg.retired_fonts:
            if font.lower() in lowered:
                findings.append(
                    Finding("error", rel, lineno, "retired-font", font)
                )

        if banned_hexes:
            hits: set[str] = set()
            for match in HEX_WITH_HASH_RE.finditer(line):
                value = match.group(1).lower()
                value = "".join(c * 2 for c in value) if len(value) == 3 else value
                if value in banned_hexes:
                    hits.add(f"#{value}")
            for match in HEX_QUOTED_RE.finditer(line):
                if match.group(1).lower() in banned_hexes:
                    hits.add(f"#{match.group(1).lower()}")
            for hit in sorted(hits):
                findings.append(
                    Finding("error", rel, lineno, "banned-hex", hit)
                )

        if cfg.forbid_em_dash and EM_DASH in line:
            findings.append(
                Finding("error", rel, lineno, "em-dash", "U+2014 em dash present")
            )
        if cfg.forbid_en_dash and EN_DASH in line:
            findings.append(
                Finding("error", rel, lineno, "en-dash", "U+2013 en dash present")
            )

    return findings


def extract_docx_text(path: Path) -> str | None:
    """Best-effort text extraction from .docx (requires python-docx)."""
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError:
        print(
            f"warning: python-docx not installed; skipping {path}", file=sys.stderr
        )
        return None
    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001 - library raises varied types
        print(f"warning: cannot open {path}: {exc}", file=sys.stderr)
        return None
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Brand compliance lint gate.")
    parser.add_argument("config", type=Path, help="brand config JSON")
    parser.add_argument("paths", nargs="+", type=Path, help="files or directories")
    parser.add_argument(
        "--ext",
        help="comma-separated extension override (e.g. .md,.html)",
    )
    parser.add_argument("--json", action="store_true", help="JSON output")
    parser.add_argument(
        "--warnings-as-errors", action="store_true", help="warnings fail the gate"
    )
    parser.add_argument(
        "--docx", action="store_true", help="also lint .docx text (needs python-docx)"
    )
    args = parser.parse_args(argv)

    cfg = BrandConfig.load(args.config)

    extensions = DEFAULT_EXTENSIONS
    if args.ext:
        extensions = frozenset(
            e if e.startswith(".") else f".{e}"
            for e in (s.strip().lower() for s in args.ext.split(","))
            if e
        )
        if not extensions:
            print("config error: --ext produced no extensions", file=sys.stderr)
            return 2
    if args.docx:
        extensions = frozenset(extensions | {".docx"})

    findings: list[Finding] = []
    files_scanned = 0
    for path in iter_target_files(list(args.paths), extensions):
        try:
            if path.stat().st_size > MAX_FILE_BYTES:
                print(f"warning: skipping oversized file {path}", file=sys.stderr)
                continue
            if path.suffix.lower() == ".docx":
                text = extract_docx_text(path)
                if text is None:
                    continue
            else:
                text = path.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            print(f"warning: cannot read {path}: {exc}", file=sys.stderr)
            continue
        files_scanned += 1
        findings.extend(lint_text(path, text, cfg))

    errors = [f for f in findings if f.severity == "error"]
    warnings = [f for f in findings if f.severity == "warning"]

    if args.json:
        print(
            json.dumps(
                {
                    "brand": cfg.brand,
                    "version": cfg.version,
                    "files_scanned": files_scanned,
                    "errors": len(errors),
                    "warnings": len(warnings),
                    "findings": [f.to_dict() for f in findings],
                },
                indent=2,
            )
        )
    else:
        print(f"brand_lint: {cfg.brand} v{cfg.version} | files scanned: {files_scanned}")
        for f in findings:
            print(f"{f.severity.upper():7} {f.path}:{f.line} [{f.rule}] {f.detail}")
        print(f"summary: {len(errors)} error(s), {len(warnings)} warning(s)")

    if files_scanned == 0:
        print("warning: no matching files scanned", file=sys.stderr)

    if errors or (args.warnings_as_errors and warnings):
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
